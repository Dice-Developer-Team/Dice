# -*- coding: utf-8 -*-
from random import randint
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import re

namedict = {}

def clear(sen):
    new = ""
    for t in sen:
        if t == "":
            continue
        new += t
    return new

def clearSent(sen):
    clr = ""
    flag = 0
    pic = 0
    for t in sen:
        if sen[0] == u"(" or sen[0] == u"（":
            flag += 1
        if sen[0] == u")" or sen[0] == u"）":
            flag -= 1
        if flag == 0 and pic == 0:
            clr += t
    # 只能去一次图片
    b = clr.find("[CQ:image")
    e = clr.find("]", b, len(clr))
    nclr = ''
    if b != -1 and e != -1:
        nclr = clr[0:b] + clr[e+1:]
    else:
        nclr = clr
    return nclr

def difColor(color):
    maxDif = 0
    for oneColor in namedict.itervalues():
        dif = sum(map(abs, [abs(color[_] - oneColor[_]) for _ in range(3)]))
        if maxDif < dif:
            maxDif = dif
    return maxDif

def ranColor():
    color = [randint(0, 255) for _ in range(3)]
    color_sum = sum(color)
    tryTime = 0
    while((color_sum < 300 or difColor(color) <50) and tryTime<100):
        color = [randint(0, 255) for _ in range(3)]

def gen(list):
    savepath = list[0]
    document = Document()
    document.styles['Normal'].font.name=u'微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    color = [randint(0, 300) for _ in range(3)]
    namedict["@@@@@section@@@@@"] = color
    for sen in list[1:]:
        sen = clear(sen)
        try:
            sen = sen.decode('GBK')
        except:
            try:
                sen = sen.decode('utf8')
            except:
                try:
                    sen = sen.decode('GB2312')
                except:
                    pass

        name, action = sen.split(":",1)
        if name == u"【下列log所属场景】":
            name = "@@@@@section@@@@@"
            style = document.styles['Normal']
            font = style.font
            font.size = Pt(20)
            namedict[name] = [randint(0, 255) for _ in range(3)]
            run = document.add_heading("", level=0).add_run(action)
            run.font.color.rgb = RGBColor(*namedict[name])
            continue
        else:
            if action[0] == u".":
                continue
            if action == u"#":
                continue
            if namedict.has_key(name):
                color = namedict[name]
            else:
                color = [randint(0, 255) for _ in range(3)]
                namedict[name] = color
            newaction = clearSent(action)
            if len(newaction) == 0:
                continue
            p = document.add_paragraph()
            style = document.styles['Normal']
            font = style.font
            font.size = Pt(8)
            if action[0] == u"#":
                if newaction[1:].find(name) == 0:
                    run = p.add_run(newaction[1:])
                else:
                    run = p.add_run(name + newaction[1:])
            else:
                if newaction.find(name) == 0:
                    run = p.add_run(name + u"：" + newaction[len(name):])
                else:
                    run = p.add_run(name + u"：" + newaction)
            run.font.color.rgb = RGBColor(*color)

    document.save(savepath)
    with open(savepath + '.finish', 'w') as fflag:
        fflag.write("finish")


def genFromTxt(infile):
    with open(infile, 'r') as fin:
        list = fin.readlines()
    document = Document()
    document.add_heading("xxx", 0)

    for x_list in list[0:]:
        p = document.add_paragraph()
        xxxx = x_list.decode('GBK')
        run = p.add_run(xxxx)
        color = (randint(0, 255) for _ in range(3))
        run.font.color.rgb = RGBColor(*color)

    document.save(infile + '.docx')
